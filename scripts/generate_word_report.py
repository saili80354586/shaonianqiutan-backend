#!/usr/bin/env python3
"""
Generate Word report from MD files for the 少年球探 platform.
Usage: python3 generate_word_report.py <rating_md_path> <player_info_md_path> <output_dir> <player_name>
"""

import sys
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def heading_style(paragraph, level=1):
    """Apply heading style to paragraph"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading(doc, text, level=1):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 100, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 80, 80)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with proper styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def parse_md_content(md_text):
    """Parse MD content and return structured data"""
    sections = []
    current_section = {'title': '', 'content': []}

    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Heading 1
        if line.startswith('# '):
            if current_section['title'] or current_section['content']:
                sections.append(current_section)
            current_section = {'title': line[2:].strip(), 'level': 1, 'content': []}
        # Heading 2
        elif line.startswith('## '):
            if current_section['title'] or current_section['content']:
                sections.append(current_section)
            current_section = {'title': line[3:].strip(), 'level': 2, 'content': []}
        # Heading 3
        elif line.startswith('### '):
            if current_section['title'] or current_section['content']:
                sections.append(current_section)
            current_section = {'title': line[4:].strip(), 'level': 3, 'content': []}
        # List item
        elif line.startswith('- '):
            current_section['content'].append(('list', line[2:]))
        # Table row
        elif line.startswith('|'):
            current_section['content'].append(('table', line))
        else:
            # Handle bold text
            cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            cleaned = re.sub(r'<br>', '\n', cleaned)
            if cleaned:
                current_section['content'].append(('text', cleaned))

    if current_section['title'] or current_section['content']:
        sections.append(current_section)

    return sections

def create_report_summary(sections):
    """Create a text summary from all sections for the Word document"""
    summary = []
    for section in sections:
        if section['level'] == 1:
            summary.append(f"\n{'='*60}\n{section['title']}\n{'='*60}")
        elif section['level'] == 2:
            summary.append(f"\n{'-'*40}\n{section['title']}\n{'-'*40}")
        elif section['level'] == 3:
            summary.append(f"\n>> {section['title']}")

        for item_type, item_content in section['content']:
            if item_type == 'text':
                summary.append(item_content)
            elif item_type == 'list':
                summary.append(f"  • {item_content}")
            elif item_type == 'table':
                summary.append(f"  {item_content}")

    return '\n'.join(summary)

def generate_word_report(rating_md_path, player_info_md_path, output_dir, player_name):
    """Generate Word report from MD files"""
    # Read MD files
    with open(rating_md_path, 'r', encoding='utf-8') as f:
        rating_content = f.read()

    with open(player_info_md_path, 'r', encoding='utf-8') as f:
        player_info_content = f.read()

    # Parse MD files
    rating_sections = parse_md_content(rating_content)
    player_sections = parse_md_content(player_info_content)

    # Create Word document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading(f'{player_name} - 比赛视频分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 100, 0)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ========== Part 1: 球员基础信息 ==========
    add_heading(doc, "第一部分：球员基础信息", 1)

    # Extract player info from the parsed content
    for section in player_sections:
        if '球员基础信息' in section['title']:
            for item_type, item_content in section['content']:
                if item_type == 'text':
                    # Parse key-value pairs like "**姓名：** 王小明<br>"
                    kv_match = re.match(r'\*\*([^：]+)：\*\*\s*(.+)', item_content)
                    if kv_match:
                        key = kv_match.group(1)
                        value = kv_match.group(2).replace('<br>', '').strip()
                        p = doc.add_paragraph()
                        run = p.add_run(f"{key}：")
                        run.bold = True
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        run = p.add_run(value)
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        elif '足球经历' in section['title'] or '技术标签' in section['title'] or '比赛风格' in section['title']:
            add_heading(doc, section['title'], 2)
            for item_type, item_content in section['content']:
                if item_type == 'text':
                    p = doc.add_paragraph()
                    run = p.add_run(item_content.replace('<br>', '\n'))
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                elif item_type == 'list':
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(item_content)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Add page break
    doc.add_page_break()

    # ========== Part 2: 球员评分报告 ==========
    add_heading(doc, "第二部分：球员评分报告", 1)

    for section in rating_sections:
        if section['level'] == 1 and '球员评分报告' in section['title']:
            # Skip the main title, process content
            continue

        if section['level'] >= 2:
            add_heading(doc, section['title'], min(section['level'], 3))

        for item_type, item_content in section['content']:
            if item_type == 'text':
                # Handle bold markers and HTML breaks
                cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', item_content)
                cleaned = re.sub(r'<br>', '\n', cleaned)

                if cleaned.startswith('**') and '：**' in cleaned:
                    # Key-value format
                    p = doc.add_paragraph()
                    parts = cleaned.split('：', 1)
                    if len(parts) == 2:
                        run = p.add_run(parts[0] + '：')
                        run.bold = True
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        run = p.add_run(parts[1])
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(cleaned)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif item_type == 'list':
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item_content)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif item_type == 'table':
                # Simple table handling - skip for now as complex tables need special handling
                pass

    # Save document
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    safe_player_name = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '_', player_name)
    output_filename = f"{safe_player_name}_比赛分析报告_{timestamp}.docx"
    output_path = os.path.join(output_dir, output_filename)

    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    if len(sys.argv) < 5:
        print("Usage: python3 generate_word_report.py <rating_md_path> <player_info_md_path> <output_dir> <player_name>")
        sys.exit(1)

    rating_md_path = sys.argv[1]
    player_info_md_path = sys.argv[2]
    output_dir = sys.argv[3]
    player_name = sys.argv[4]

    try:
        output_path = generate_word_report(rating_md_path, player_info_md_path, output_dir, player_name)
        print(f"SUCCESS:{output_path}")
    except Exception as e:
        print(f"ERROR:{str(e)}")
        sys.exit(1)